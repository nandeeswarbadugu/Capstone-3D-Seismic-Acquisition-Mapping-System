 
import tkinter as tk
from tkinter import messagebox
import numpy as np
import matplotlib.pyplot as plt
from shapely.geometry import Point
from docx import Document
import os

class SeismicAcquisitionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Seismic Acquisition")

        # Set window size
        self.root.geometry("600x500")

        # Label and Entry for each parameter
        self.create_widgets()

    def create_widgets(self):
        # Survey Extent Inputs
        self.label_x_extent = tk.Label(self.root, text="Survey Length (X):")
        self.label_x_extent.grid(row=0, column=0, padx=10, pady=5)
        self.entry_x_extent = tk.Entry(self.root)
        self.entry_x_extent.grid(row=0, column=1, padx=10, pady=5)

        self.label_y_extent = tk.Label(self.root, text="Survey Width (Y):")
        self.label_y_extent.grid(row=1, column=0, padx=10, pady=5)
        self.entry_y_extent = tk.Entry(self.root)
        self.entry_y_extent.grid(row=1, column=1, padx=10, pady=5)

        # Geophone and Source Distance Inputs
        self.label_gx_distance = tk.Label(self.root, text="Geophone Distance:")
        self.label_gx_distance.grid(row=2, column=0, padx=10, pady=5)
        self.entry_gx_distance = tk.Entry(self.root)
        self.entry_gx_distance.grid(row=2, column=1, padx=10, pady=5)

        self.label_gx_line_distance = tk.Label(self.root, text="Geophone Line Distance:")
        self.label_gx_line_distance.grid(row=3, column=0, padx=10, pady=5)
        self.entry_gx_line_distance = tk.Entry(self.root)
        self.entry_gx_line_distance.grid(row=3, column=1, padx=10, pady=5)

        self.label_sx_distance = tk.Label(self.root, text="Source Distance:")
        self.label_sx_distance.grid(row=4, column=0, padx=10, pady=5)
        self.entry_sx_distance = tk.Entry(self.root)
        self.entry_sx_distance.grid(row=4, column=1, padx=10, pady=5)

        self.label_sx_line_distance = tk.Label(self.root, text="Source Line Distance:")
        self.label_sx_line_distance.grid(row=5, column=0, padx=10, pady=5)
        self.entry_sx_line_distance = tk.Entry(self.root)
        self.entry_sx_line_distance.grid(row=5, column=1, padx=10, pady=5)

        # Buttons for actions
        self.plot_button = tk.Button(self.root, text="Plot Seismic Data", command=self.plot_seismic_data)
        self.plot_button.grid(row=6, column=0, columnspan=2, pady=10)

        self.save_button = tk.Button(self.root, text="Save Multiplicity to File", command=self.save_multiplicity)
        self.save_button.grid(row=7, column=0, columnspan=2, pady=10)

    def get_valid_input(self, value):
        try:
            value = float(value)
            if value <= 0:
                raise ValueError
            return value
        except ValueError:
            messagebox.showerror("Invalid Input", "Please enter a non-zero positive numeric value.")
            return None

    def plot_seismic_data(self):
        # Get values from the entry widgets
        x_extent = self.get_valid_input(self.entry_x_extent.get())
        y_extent = self.get_valid_input(self.entry_y_extent.get())
        gx_distance = self.get_valid_input(self.entry_gx_distance.get())
        gx_line_distance = self.get_valid_input(self.entry_gx_line_distance.get())
        sx_distance = self.get_valid_input(self.entry_sx_distance.get())
        sx_line_distance = self.get_valid_input(self.entry_sx_line_distance.get())

        if None in [x_extent, y_extent, gx_distance, gx_line_distance, sx_distance, sx_line_distance]:
            return  # Exit if any input is invalid

        # Adjust starting points based on input distances
        x_start_geophones = gx_line_distance
        y_start_geophones = gx_distance
        x_start_sources = sx_line_distance
        y_start_sources = sx_distance

        # Ensure geophones and sources do not exceed the survey area
        geophone_lines = int((y_extent - y_start_geophones) / gx_line_distance) + 1
        geophones_per_line = int((x_extent - x_start_geophones) / gx_distance) + 1
        source_lines = int((x_extent - x_start_sources) / sx_line_distance) + 1
        sources_per_line = int((y_extent - y_start_sources) / sx_distance) + 1

# Generate geophone positions
        rcvrx, rcvry = zip(*[
            (x_start_geophones + i * gx_distance, y_start_geophones + j * gx_line_distance)
            for j in range(geophone_lines)
            for i in range(geophones_per_line)
            if y_start_geophones + j * gx_line_distance <= y_extent  # Ensure geophone line within survey area
        ])

# Generate source positions
        srcx, srcy = zip(*[
            (x_start_sources + i * sx_line_distance, y_start_sources + j * sx_distance)
            for i in range(source_lines)
            for j in range(sources_per_line)
        ])

        geophones = [Point(x, y) for x, y in zip(rcvrx, rcvry)]
        sources = [Point(x, y) for x, y in zip(srcx, srcy)]

        # Calculate midpoints, offsets, and multiplicity
        midpoints = []
        multiplicity = {}
        offsets = []
        for geophone in geophones:
            for source in sources:
                midpoint_x = (geophone.x + source.x) / 2
                midpoint_y = (geophone.y + source.y) / 2
                midpoint = Point(midpoint_x, midpoint_y)
                midpoints.append(midpoint)

                # Calculate offset
                offset = source.distance(geophone)
                offsets.append(offset)

                # Update multiplicity
                if (midpoint_x, midpoint_y) in multiplicity:
                    multiplicity[(midpoint_x, midpoint_y)] += 1
                else:
                    multiplicity[(midpoint_x, midpoint_y)] = 1

        # Plot geophones, sources, and midpoints
        plt.figure(figsize=(14, 8))
        plt.scatter([g.x for g in geophones], [g.y for g in geophones], c='blue', label='Geophones', marker='^', s=100)
        plt.scatter([s.x for s in sources], [s.y for s in sources], c='red', label='Sources', marker='o', s=100)
        plt.scatter([m.x for m in midpoints], [m.y for m in midpoints], c='green', label='Bins', marker='o', s=30, alpha=0.6)

# Annotate geophones
        for x, y in zip([g.x for g in geophones], [g.y for g in geophones]):
            plt.text(x-1, y-1, "G", fontsize=8, color='blue', ha='right')

# Annotate sources
        for x, y in zip([s.x for s in sources], [s.y for s in sources]):
            plt.text(x+1, y+1, "S", fontsize=8, color='red', ha='left')
            
            # Annotate midpoints with multiplicity
        for m in multiplicity:
            x, y = m
            plt.text(x+1.5, y+1.5, f"{multiplicity[m]}", fontsize=8, color='green', ha='center')

        plt.title("Seismic Acquisition - Geophones, Sources, and Bins")
        plt.xlabel("Area (length)")
        plt.ylabel("Area (width)")
        plt.grid(True)
        plt.legend()
        plt.show()

        # Plot offset distribution
        plt.figure(figsize=(10, 6))
        plt.plot(offsets, color='blue', marker='o', linestyle='-', alpha=0.7)
        plt.title("Offset Distribution")
        plt.xlabel("Index")
        plt.ylabel("Offset (m)")
        plt.grid(True)
        plt.show()

    def save_multiplicity(self):
        # Get values from the entry widgets
        x_extent = self.get_valid_input(self.entry_x_extent.get())
        y_extent = self.get_valid_input(self.entry_y_extent.get())
        gx_distance = self.get_valid_input(self.entry_gx_distance.get())
        gx_line_distance = self.get_valid_input(self.entry_gx_line_distance.get())
        sx_distance = self.get_valid_input(self.entry_sx_distance.get())
        sx_line_distance = self.get_valid_input(self.entry_sx_line_distance.get())

        if None in [x_extent, y_extent, gx_distance, gx_line_distance, sx_distance, sx_line_distance]:
            return  # Exit if any input is invalid

        # Adjust starting points based on input distances
        x_start_geophones = gx_line_distance
        y_start_geophones = gx_distance
        x_start_sources = sx_line_distance
        y_start_sources = sx_distance

        # Ensure geophones and sources do not exceed the survey area
        geophone_lines = int((y_extent - y_start_geophones) / gx_line_distance) + 1
        geophones_per_line = int((x_extent - x_start_geophones) / gx_distance) + 1
        source_lines = int((x_extent - x_start_sources) / sx_line_distance) + 1
        sources_per_line = int((y_extent - y_start_sources) / sx_distance) + 1

        # Generate geophone positions
        rcvrx, rcvry = zip(*[
            (x_start_geophones + i * gx_distance, y_start_geophones + j * gx_line_distance)
            for j in range(geophone_lines)
            for i in range(geophones_per_line)
            if y_start_geophones + j * gx_line_distance <= y_extent  # Ensure geophone line within survey area
        ])

        # Generate source positions
        srcx, srcy = zip(*[
            (x_start_sources + i * sx_line_distance, y_start_sources + j * sx_distance)
            for i in range(source_lines)
            for j in range(sources_per_line)
        ])

        geophones = [Point(x, y) for x, y in zip(rcvrx, rcvry)]
        sources = [Point(x, y) for x, y in zip(srcx, srcy)]

        # Calculate midpoints, offsets, and multiplicity
        midpoints = []
        multiplicity = {}
        offsets = []

        for geophone in geophones:
            for source in sources:
                midpoint_x = (geophone.x + source.x) / 2
                midpoint_y = (geophone.y + source.y) / 2
                midpoint = Point(midpoint_x, midpoint_y)
                midpoints.append(midpoint)

                # Calculate offset
                offset = geophone.distance(source)
                offsets.append(offset)

                # Update multiplicity
                if (midpoint_x, midpoint_y) in multiplicity:
                    multiplicity[(midpoint_x, midpoint_y)] += 1
                else:
                    multiplicity[(midpoint_x, midpoint_y)] = 1

        # Save multiplicity to document
        multiplicity_by_line = {}
        for (x, y), count in multiplicity.items():
            if y not in multiplicity_by_line:
                multiplicity_by_line[y] = []
            multiplicity_by_line[y].append(count)

        desktop_path = os.path.join(os.path.expanduser("~"), "downloads")
        file_name = "multiplicity_values.docx"
        file_path = os.path.join(desktop_path, file_name)

        doc = Document()
        doc.add_heading("Multiplicity Values", 0)

        try:
            for line_y in sorted(multiplicity_by_line.keys()):
                line_values = multiplicity_by_line[line_y]
                doc.add_paragraph("  " + " ".join(map(str, line_values)))

            doc.save(file_path)
            messagebox.showinfo("Success", f"Multiplicity values have been saved to {file_path}")

 # Open the file after saving
            if os.name == 'nt':  # For Windows
                os.startfile(file_path)
            elif os.name == 'posix':  # For macOS or Linux
                subprocess.call(['open', file_path])  # macOS
            # subprocess.call(['xdg-open', file_path])  # Linux
            else:
                messagebox.showwarning("Unsupported OS", "Unable to open the file automatically on this operating system.")

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while saving the file: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = SeismicAcquisitionApp(root)
    root.mainloop()

